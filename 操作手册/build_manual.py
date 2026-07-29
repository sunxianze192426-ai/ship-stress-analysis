#!/usr/bin/env python3
"""
构建操作手册 docx —— 完整复制原文档风格 + 新增5个V2.0功能章节
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 路径
# ============================================================
OUT_DIR = r"C:\Users\Xianze Sun\Desktop\demo1\操作手册"
IMG_DIR = os.path.join(OUT_DIR, "images")
OUT_PATH = os.path.join(OUT_DIR, "操作手册（舰船平台综合环境应力剖面分析系统）V2.0.docx")

doc = Document()

# ============================================================
# 0. 页面设置 + 样式
# ============================================================
sec = doc.sections[0]
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin   = Cm(3.17)
sec.right_margin  = Cm(3.17)
sec.top_margin    = Cm(2.54)
sec.bottom_margin = Cm(2.54)

# ── Normal 样式：两端对齐 ──
style_normal = doc.styles['Normal']
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
# 字号/字体不设，由 run 级控制

# ============================================================
# 辅助函数
# ============================================================
def set_line_spacing_1_5(para):
    """1.5倍行距 (line=360, lineRule=auto)"""
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'auto')

def set_first_line_indent(para, cm_val=0.85):
    """首行缩进 (480 twips = 0.85cm)"""
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), '480')

def set_left_indent(para, cm_val=0.85):
    """左缩进 (480 twips = 0.85cm)"""
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), '480')

def set_left_and_first_indent(para, left_cm=0.85, first_cm=0.85):
    """同时设置左缩进和首行缩进"""
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), '480')
    ind.set(qn('w:firstLine'), '480')

def set_font(run, name='Times New Roman', ea_name=None, size_pt=None, bold=None):
    """设置 run 的字体"""
    run.font.name = name
    if ea_name:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), ea_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold

def add_paragraph_with_text(doc, text, font_name='Times New Roman', ea_name=None,
                             size_pt=12, bold=False, alignment=None,
                             first_indent=False, left_indent=0,
                             line_1_5=True):
    """添加段落并设置文字"""
    para = doc.add_paragraph()
    if alignment is not None:
        para.paragraph_format.alignment = alignment
    if first_indent:
        set_first_line_indent(para)
    if left_indent > 0:
        set_left_indent(para, left_indent)
    if line_1_5:
        set_line_spacing_1_5(para)

    run = para.add_run(text)
    set_font(run, name=font_name, ea_name=ea_name, size_pt=size_pt, bold=bold)
    return para

def add_cover_title_para(doc, text, ea_name='黑体', size_pt=26, bold=True):
    """封面大字段落（居中，无行距设置）"""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_font(run, ea_name=ea_name, size_pt=size_pt, bold=bold)
    return para

def add_label_value_para(doc, label, value, left_indent_val=0, first_line_val=0.85):
    """封面信息区的 标签(加粗) + 值(不加粗) 段落"""
    para = doc.add_paragraph()
    set_line_spacing_1_5(para)
    if left_indent_val > 0:
        set_left_indent(para, left_indent_val)
    if first_line_val > 0:
        set_first_line_indent(para, first_line_val)

    run_label = para.add_run(label)
    set_font(run_label, size_pt=12, bold=True)
    run_val = para.add_run(value)
    set_font(run_val, size_pt=12, bold=False)
    return para

def add_section_title(doc, text, size_pt=14, bold=True):
    """章节标题：14pt 加粗，无缩进，1.5倍行距"""
    return add_paragraph_with_text(doc, text, size_pt=size_pt, bold=bold,
                                    first_indent=False, line_1_5=True)

def add_body_text(doc, text):
    """正文：12pt，首行缩进0.85cm，1.5倍行距，两端对齐"""
    return add_paragraph_with_text(doc, text, size_pt=12, bold=False,
                                    first_indent=True, line_1_5=True)

def add_figure_caption(doc, text):
    """图表标题：10.5pt，居中，不加粗，1.5倍行距"""
    return add_paragraph_with_text(doc, text, size_pt=10.5, bold=False,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    first_indent=False, line_1_5=True)

def add_image_placeholder(doc, img_name, width_cm=14.65, height_cm=8.2):
    """插入图片（居中），前后各留空行"""
    # 空行
    para_empty1 = doc.add_paragraph()
    para_empty1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing_1_5(para_empty1)

    # 图片
    img_path = os.path.join(IMG_DIR, img_name)
    para_img = doc.add_paragraph()
    para_img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing_1_5(para_img)
    if os.path.exists(img_path):
        run_img = para_img.add_run()
        run_img.add_picture(img_path, width=Cm(width_cm), height=Cm(height_cm))
    else:
        # 占位文字
        run_img = para_img.add_run(f"【截图位置：{img_name}】")
        set_font(run_img, size_pt=10.5, bold=False)

    # 空行
    para_empty2 = doc.add_paragraph()
    para_empty2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing_1_5(para_empty2)

    return para_img

def add_empty_centered(doc):
    """添加居中空行（用于图片前后留白）"""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing_1_5(para)
    return para


# ============================================================
#  生成封面占位图（用于新增模块截图位置）
# ============================================================
def create_placeholder_image(path, text, width_px=1200, height_px=700):
    """用 PIL 创建占位截图"""
    try:
        from PIL import Image, ImageDraw, ImageFont
        img = Image.new('RGB', (width_px, height_px), color=(245, 245, 245))
        draw = ImageDraw.Draw(img)
        # 边框
        draw.rectangle([0, 0, width_px-1, height_px-1], outline=(180, 180, 180), width=2)
        # 虚线效果
        for i in range(4, width_px-4, 16):
            draw.line([(i, 0), (i+8, 0)], fill=(180,180,180), width=2)
            draw.line([(i, height_px-1), (i+8, height_px-1)], fill=(180,180,180), width=2)
        for i in range(4, height_px-4, 16):
            draw.line([(0, i), (0, i+8)], fill=(180,180,180), width=2)
            draw.line([(width_px-1, i), (width_px-1, i+8)], fill=(180,180,180), width=2)

        # 文字
        try:
            font = ImageFont.truetype("simhei.ttf", 32)
            font_small = ImageFont.truetype("simhei.ttf", 22)
        except:
            font = ImageFont.load_default()
            font_small = font

        bbox = draw.textbbox((0,0), text, font=font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        draw.text(((width_px-tw)//2, (height_px-th)//2 - 20), text, fill=(120,120,120), font=font)

        hint = "请在此处插入截图"
        bbox2 = draw.textbbox((0,0), hint, font=font_small)
        tw2 = bbox2[2] - bbox2[0]
        draw.text(((width_px-tw2)//2, (height_px-th)//2 + 30), hint, fill=(160,160,160), font=font_small)

        img.save(path)
        print(f"Created placeholder: {path}")
    except Exception as e:
        print(f"PIL not available, skipping placeholder ({e})")


# 生成新增模块的占位图
for i, label in enumerate([
    "图21 设备应力综合分析模块主界面",
    "图22 蒸汽压力分析计算结果",
    "图23 蒸汽流量分析计算结果",
    "图24 蒸汽颗粒物浓度分析计算结果",
    "图25 盐雾浓度分析计算结果",
    "图26 氧浓度分析计算结果",
], start=21):
    create_placeholder_image(
        os.path.join(IMG_DIR, f"placeholder_fig{i}.png"),
        label
    )


print("=" * 60)
print("开始构建文档...")
print("=" * 60)

# ============================================================
# 第一部分：封面
# ============================================================

# 6个居中空行
for _ in range(6):
    add_cover_title_para(doc, '', size_pt=12, bold=False)

# 系统名称 — 22pt 黑体 加粗 居中
add_cover_title_para(doc, '舰船平台综合环境应力剖面分析系统', ea_name='黑体', size_pt=22)

# "操作手册"四个大字 — 26pt 黑体 加粗 居中
for ch in ['操', '作', '手', '册']:
    add_cover_title_para(doc, ch, ea_name='黑体', size_pt=26)

# 4个居中空行
for _ in range(4):
    add_cover_title_para(doc, '', size_pt=12, bold=False)

# 单位 — 16pt 仿宋 加粗 居中
add_cover_title_para(doc, '浙江理工大学', ea_name='仿宋', size_pt=16)

# 日期 — 16pt 仿宋 加粗 居中
add_cover_title_para(doc, '二零二三年十一月', ea_name='仿宋', size_pt=16)

# 1个居中空行
add_cover_title_para(doc, '', size_pt=12, bold=False)

# ── 软件基本信息区（1.5倍行距）──
add_label_value_para(doc, '软件名称：', '舰船平台综合环境应力剖面分析系统 V2.0',
                     left_indent_val=0, first_line_val=0)
add_label_value_para(doc, '完成日期：', '2026年7月29日',
                     left_indent_val=0, first_line_val=0)
add_label_value_para(doc, '软件版本号：', 'V2.0',
                     left_indent_val=0, first_line_val=0)

# 运行环境（加粗标签）
p = add_label_value_para(doc, '运行环境：', '', left_indent_val=0, first_line_val=0)

# 硬件环境（左缩进0.85cm）
p = doc.add_paragraph()
set_line_spacing_1_5(p)
set_left_indent(p, 0.85)
r = p.add_run('硬件环境：')
set_font(r, size_pt=12, bold=False)

# 硬件环境具体项（左缩进 + 首行缩进）
for item_text in [
    '开发软件机型及CPU型号：CORE i5',
    '主要适应机型及CPU型号：CORE i3及以上',
    '内存要求：8.00GB以上',
    '终端要求：无',
]:
    add_label_value_para(doc, item_text, '',
                         left_indent_val=0.85, first_line_val=0.85)

# 软件环境
p = doc.add_paragraph()
set_line_spacing_1_5(p)
set_left_indent(p, 0.85)
r = p.add_run('软件环境：')
set_font(r, size_pt=12, bold=False)

# 软件环境具体项
for item_text in [
    '运行操作系统名称及版本号：Windows 10或Windows 11',
    '编程语言名称及版本号：MATLAB R2021a',
    '源程序量：约3500行（V2.0新增约1127行）',
    '程序存储媒体：硬盘',
]:
    add_label_value_para(doc, item_text, '',
                         left_indent_val=0.85, first_line_val=0.85)

# 软件说明
add_label_value_para(doc, '软件说明：', '', left_indent_val=0, first_line_val=0)

add_body_text(doc,
    '该软件主要用于对舰船平台及装备的综合环境应力（包括振动、冲击、温湿度、倾角等）'
    '进行数据处理及分析，得到试验剖面。该软件首先通过在登录界面右侧输入正确的账号及密码，'
    '点击按钮"登录"，即可跳转主窗口界面。然后可以点击主界面顶部选择不同的剖面分析模块，'
    '分别为振动剖面模块、冲击剖面模块、温湿度剖面1模块、温湿度剖面2模块及倾角剖面模块。'
    '在每个模块左上角可以选择数据文件位置、输入数据相关参数（数据位置、长度、采样频率等），'
    '点击按钮生成时域图、分析图及试验谱等图像，还可显示出试验谱相关数据。'
    'V2.0版本新增了第6个标签页——设备应力综合分析模块（Tab_6），'
    '支持对5种环境应力参数（蒸汽压力、蒸汽流量、蒸汽颗粒物浓度、盐雾浓度、氧浓度）'
    '进行按设备类型分类的自动统计分析。')

add_body_text(doc,
    '本系统是基于MATLAB R2021a语言进行编程，界面明了，易操作，经过几轮更新和完善，'
    '现为V2.0版本。')

# 分页符
doc.add_page_break()

# ============================================================
# 第二部分：界面概述
# ============================================================

# 1、登录界面说明
add_section_title(doc, '1、登录界面说明')

add_body_text(doc,
    '软件的初始登录界面如图1所示，界面的右侧为账号信息输入区，'
    '用于输入账户及密码。当输入正确账户密码后，点击"登录"按钮，'
    '即可显示黑色字"登陆成功"，同时关闭登录界面，跳转主窗口界面。'
    '当输入错误账户密码后，点击"登录"按钮，可显示红色字"账号或密码错误，请重试"。')

add_image_placeholder(doc, 'image1.png', width_cm=14.2, height_cm=9.75)
add_figure_caption(doc, '图1 软件初始登录界面')

add_empty_centered(doc)

# 2、主窗口说明
add_section_title(doc, '2、主窗口说明')

add_body_text(doc,
    '软件的主窗口界面如图2所示，界面的顶部为模块选择区域，'
    '用于选择所需使用的应力剖面分析模块，主要包括振动模块、冲击模块、'
    '温湿度模块及倾角模块（V2.0新增设备应力综合分析模块）。'
    '每一个模块左上角都有一块参数输入区，用于输入数据文件所在位置、'
    '采样率以及进行分析的数据区域等。')

add_image_placeholder(doc, 'image2.png', width_cm=14.65, height_cm=8.24)
add_figure_caption(doc, '图2 软件主窗口界面')

add_empty_centered(doc)

add_body_text(doc,
    '如图2所示，振动剖面分析模块一次可输入一整组数据（不限样本数量、不限数据长度）。'
    '分别点击"画时域图"、"上限谱&规范谱"、"前期检验"部分的按钮，'
    '可分别绘制时域图、频域图、前期检验图、分析图以及试验谱等，'
    '最后显示各频段宽带谱的试验谱幅值分析结果。')

add_body_text(doc,
    '如图3所示，冲击剖面分析模块一次可输入一整组数据（不限样本数量、不限数据长度）。'
    '分别点击"画时域图"、"上限谱&规范谱"部分按钮，'
    '可分别绘制时域图、频域分析图以及试验谱等，最后显示各频段试验谱幅值分析结果。')

add_image_placeholder(doc, 'image3.png', width_cm=14.65, height_cm=8.24)
add_figure_caption(doc, '图3 冲击剖面分析模块')

add_empty_centered(doc)

add_body_text(doc,
    '如图4、5所示，温湿度剖面分析模块可以同时输入一组温度数据和一组湿度数据。'
    '点击"剖面计算"按钮，可分别绘制温度和湿度的时域图、分析图、试验剖面图等，'
    '最后显示试验剖面相关数据分析结果。')

add_image_placeholder(doc, 'image4.png', width_cm=14.65, height_cm=8.24)
add_figure_caption(doc, '图4 温湿度剖面分析模块-前期处理部分')

add_empty_centered(doc)

add_image_placeholder(doc, 'image5.png', width_cm=14.65, height_cm=8.15)
add_figure_caption(doc, '图5 温湿度剖面分析模块-剖面分析部分')

add_empty_centered(doc)

add_body_text(doc,
    '如图6所示，倾角分析模块可以同时输入三组三个轴向的倾角数据。'
    '分别点击"x轴计算"、"y轴计算"、"z轴计算"部分的按钮，'
    '可以分别绘制各轴向倾角的时域图、倾角变化-时间关系离散图及试验谱等，'
    '最后分别显示各轴向试验谱的幅值与周期值。')

add_image_placeholder(doc, 'image6.png', width_cm=14.65, height_cm=8.15)
add_figure_caption(doc, '图6 倾角剖面分析模块')

add_empty_centered(doc)

# 2、功能使用
add_section_title(doc, '2、功能使用')

# ============================================================
# 2.1 登录界面
# ============================================================
add_section_title(doc, '2.1 登录界面')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，如图7所示，首先在账号登录区对应方框内输入本次登录软件的账户和密码，'
    '然后点击"登录"按钮，程序就会进行账号核对。')

add_image_placeholder(doc, 'image7.png', width_cm=4.8, height_cm=5.83)
add_figure_caption(doc, '图7 账号登录区')

add_empty_centered(doc)

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '当输入的账户及密码是正确的时，按钮上方即显示黑色字"登陆成功"，'
    '同时关闭登录界面，跳转至主窗口界面。当输入的账户及密码是错误的时，'
    '即显示红色字"账号或密码错误，请重试"，此时不会自动进行下一步操作，'
    '直至账户和密码输入正确。')

add_empty_centered(doc)

# ============================================================
# 2.2 振动剖面分析模块
# ============================================================
add_section_title(doc, '2.2 振动剖面分析模块')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，如图8所示，首先点击"选择数据"按钮即可在电脑硬盘中选择数据文件，'
    '然后在对应方框中输入我们分析的数据在文件中的列数、数据上下限、分段数、'
    '数据采样率以及频域分析的谱线数。')

add_image_placeholder(doc, 'image8.png', width_cm=14.65, height_cm=3.05)
add_figure_caption(doc, '图8 振动部分参数输入区')

add_empty_centered(doc)

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '通过点击图8中前期检验的正态性检验按钮，程序开始运行并绘制正态性检验结果图，'
    '如图9所示。同理，分别点击稳定性检验或周期性检验，会绘制稳定性检验或周期性检验结果图，'
    '如图10和图11所示。')

add_image_placeholder(doc, 'image9.png', width_cm=14.65, height_cm=9.6)
add_figure_caption(doc, '图9 正态性检验图')

add_empty_centered(doc)

add_image_placeholder(doc, 'image10.png', width_cm=14.65, height_cm=9.69)
add_figure_caption(doc, '图10 稳定性检验图')

add_empty_centered(doc)

add_image_placeholder(doc, 'image11.png', width_cm=14.65, height_cm=9.79)
add_figure_caption(doc, '图11 周期性检验图')

add_empty_centered(doc)

add_body_text(doc,
    '通过点击图8中画时域图按钮，程序开始运行并绘制图像，如图12中第一张图所示。')

add_body_text(doc,
    '通过点击图8中上限图&规范谱按钮，程序开始运行并绘制分析过程图像。'
    '如图12中第二张图，通过把所有时域图fft变换成频域图，再转换单位后绘制在一个坐标系中，'
    '分析求得实测正态单侧上限谱和规范谱。如图12中第四张图，对图二求得的上限谱进行分析，'
    '提取出大于0.01（m/ss）s/Hz的部分计算该处的窄带范围及值大小。'
    '如图12中第五张图，对图四求得的窄带分量进行剔除，剩余部分进行分段分析，'
    '对各频率段数据采用正态单侧容差上线法求得平直谱幅值，估计参数选取P99/90，'
    '得到宽带分量的规范谱。如图12中第六张图，对图四和图五求得的窄带规范谱和宽带规范谱'
    '进行叠加得到最终试验谱，程序运行完毕。')

add_image_placeholder(doc, 'image12.png', width_cm=14.65, height_cm=8.15)
add_figure_caption(doc, '图12 振动剖面分析计算结果')

add_empty_centered(doc)

add_body_text(doc,
    '如图12右上角所示，程序运行结束后，会在界面的右上角显示出宽带试验谱各频率区间内的'
    '振动功率谱密度值。')

add_empty_centered(doc)

# ============================================================
# 2.3 冲击剖面分析模块
# ============================================================
add_section_title(doc, '2.3 冲击剖面分析模块')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，如图13所示，首先点击"选择数据"按钮即可在电脑硬盘中选择数据文件，'
    '然后在对应方框中输入我们分析的数据在文件中的列数、数据上下限、列数范围以及数据采样率等。')

add_image_placeholder(doc, 'image13.png', width_cm=14.65, height_cm=5.73)
add_figure_caption(doc, '图13 冲击部分参数输入区')

add_empty_centered(doc)

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '通过点击图13中画时域图按钮，程序开始运行并绘制时域图，如图14所示。')

add_image_placeholder(doc, 'image14.png', width_cm=14.65, height_cm=8.28)
add_figure_caption(doc, '图14 冲击时域图')

add_empty_centered(doc)

add_body_text(doc,
    '通过点击图13中上限图&规范谱按钮，程序开始运行并绘制分析过程图像。'
    '如图15中右上角图，通过把所有时域图fft变换成频域图，再转换单位后绘制在一个坐标系中，'
    '通过上限包络法求得包络上限。如图15中右下角图，对求得的包络上限进行分析，'
    '对2000Hz以上数据采用正态单侧容差上线法求平直谱幅值，估计参数选取P99/90。'
    '通过对左侧输入参数的调整系数以及程序的多次计算直到得到一条满足要求的低频段规范谱。')

add_image_placeholder(doc, 'image15.png', width_cm=14.65, height_cm=8.15)
add_figure_caption(doc, '图15 冲击剖面分析计算结果')

add_empty_centered(doc)

add_body_text(doc,
    '如图15所示，程序运行结束后，会在界面的输入参数右侧显示出平直谱幅值和10Hz处幅值。')

add_empty_centered(doc)

# ============================================================
# 2.4 温湿度剖面分析模块
# ============================================================
add_section_title(doc, '2.4 温湿度剖面分析模块')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，如图16所示，首先点击"选择数据"按钮即可在电脑硬盘中选择数据文件，'
    '然后在对应方框中输入我们分析的温湿度数据在文件中的列数、数据上下限、每小时采样数以及样本天数等。')

add_image_placeholder(doc, 'image16.png', width_cm=14.65, height_cm=6.69)
add_figure_caption(doc, '图16 温湿度部分参数输入区')

add_empty_centered(doc)

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '通过点击图16中剖面计算按钮，程序开始运行并绘制时域图、分析图以及温湿度剖面结果等图像，'
    '如图17、图18所示。')

add_body_text(doc,
    '图17中左一图为温度时域数据图，左二为湿度时域数据图，'
    '中一为温度数据每天按h展开后的原始数据图像，'
    '中二为中一数据经过奇异值剔除处理、滑动平均滤波处理后得到的温度数据图，'
    '右一为中二数据按天进行平均然后通过温度疲劳损伤等效计算后得到的结果，'
    '右二为湿度数据按天平均得到的结果，右三为中二温度数据按h展开后的温度图像。')

add_image_placeholder(doc, 'image17.png', width_cm=14.65, height_cm=8.17)
add_figure_caption(doc, '图17 温湿度剖面分析-前期处理部分')

add_empty_centered(doc)

add_body_text(doc,
    '图18中，中一图为图17中前期处理结束后温度数据经过低通滤波得到的结果图，'
    '将其按天展开得到右一图。中二为图17中前期处理结束后温度数据经过高通滤波得到的结果图，'
    '将其按天展开得到右二图。')

add_image_placeholder(doc, 'image18.png', width_cm=14.65, height_cm=8.15)
add_figure_caption(doc, '图18 温湿度剖面分析-剖面分析部分')

add_empty_centered(doc)

add_body_text(doc,
    '如图18所示，程序运行结束后，会在界面的左上角显示出温度计算数据，'
    '包括循环次数、等效循环次数、热天数、冷天数、高温平均温度、低温平均温度'
    '以及高温平均湿度等结果。同时会在左下角绘制温湿度试验剖面图。')

add_empty_centered(doc)

# ============================================================
# 2.5 倾角剖面分析模块
# ============================================================
add_section_title(doc, '2.5 倾角剖面分析模块')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，如图19所示，首先点击"选择数据"按钮即可在电脑硬盘中选择数据文件，'
    '然后在对应方框中输入我们分析的倾角数据在文件中的列数、数据上下限、采样频率等。')

add_image_placeholder(doc, 'image19.png', width_cm=3.69, height_cm=8.87)
add_figure_caption(doc, '图19 倾角部分参数输入区')

add_empty_centered(doc)

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '通过点击图19中x轴计算按钮，程序开始运行并绘制x轴时域图、'
    '倾角变化-时间关系离散图及试验谱等图像。')

add_body_text(doc,
    '如图20所示，左一图为x轴向倾角时域数据图，'
    '中一图为x轴向倾角变化-时间关系离散图，右一为x轴向的试验谱。'
    '左二图为y轴向倾角时域数据图，中二图为y轴向倾角变化-时间关系离散图，'
    '右二为y轴向的试验谱。左三图为z轴向倾角时域数据图，'
    '中三图为z轴向倾角变化-时间关系离散图，右三为z轴向的试验谱。')

add_image_placeholder(doc, 'image20.png', width_cm=14.65, height_cm=8.15)
add_figure_caption(doc, '图20 倾角部分剖面分析')

add_empty_centered(doc)

add_body_text(doc,
    '如图20所示，程序运行结束后，会在界面的左下角显示出各轴向倾角计算结果数据，'
    '包括试验谱幅值、周期等。')

# ============================================================
# ★★★★★ 分页：V2.0新增内容 ★★★★★
# ============================================================
doc.add_page_break()

# ============================================================
# 3、设备应力综合分析模块（V2.0新增）
# ============================================================
add_section_title(doc, '3、设备应力综合分析模块（V2.0新增）')

add_body_text(doc,
    '设备应力综合分析模块是V2.0版本新增的第6个标签页（Tab_6），'
    '用于对船舶平台多类设备的多种环境应力参数进行自动分类统计与可视化分析。'
    '该模块支持三列格式数据文件（设备标签、数值、单位），'
    '可自动识别设备类型并按类型分组，计算各设备类型的6项统计指标'
    '（最大值、最小值、平均值、均方根值、方差、标准差），'
    '同时提供折线图（时序对比）和箱线图（分布分析）两种可视化方式。')

add_body_text(doc,
    '模块主界面如图21所示，左侧为控制区，包含5个参数模块按钮'
    '（蒸汽压力、蒸汽流量、蒸汽颗粒物浓度、盐雾浓度、氧浓度）、'
    '文件选择区、设备类型列表、Y轴范围调节和导出/刷新按钮。'
    '中部上方为折线图区（设备应力时序对比），中部下方为箱线图区（应力分布分析）。'
    '右侧为统计结果区，以表格形式展示6项统计指标。')

add_image_placeholder(doc, 'placeholder_fig21.png', width_cm=14.65, height_cm=8.2)
add_figure_caption(doc, '图21 设备应力综合分析模块主界面')

add_empty_centered(doc)

# ============================================================
# 3.1 蒸汽压力分析
# ============================================================
add_section_title(doc, '3.1 蒸汽压力分析')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，首先点击"选择数据"按钮在电脑硬盘中选择蒸汽压力数据文件'
    '（支持.xlsx/.xls/.csv/.txt格式）。数据文件须包含三列：设备标签（如"球阀1"）、'
    '数值（应力测量值）、单位（MPa）。系统会自动识别单位为"MPa"的数据行，'
    '并按设备标签中的类型名（如"球阀"）进行分组。同一设备标签可包含多条时序记录，'
    '表示该设备在不同时间点的多次测量值，系统自动汇总计算。')

add_body_text(doc,
    '在设备类型列表中，用户可按Ctrl键多选需要对比分析的设备类型，'
    '图表将仅显示选中的设备数据。Y轴范围默认为自动适配，用户也可在Y轴范围调节区'
    '手动输入最小值和最大值进行精细调节。')

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '通过点击"蒸汽压力 MPa"按钮，程序开始运行并对数据进行完整的处理流程：'
    '① 文件校验（检查路径有效性、文件存在性、格式合法性）；'
    '② readtable读取文件并自动推断格式；'
    '③ 三列解析（col1→设备标签、col2→数值、col3→单位）；'
    '④ 按标签类型名自动分类分组；'
    '⑤ 按单位匹配筛选当前模块对应数据；'
    '⑥ 对每个设备类型分别计算6项统计指标；'
    '⑦ 更新UI（结果表格、折线图、箱线图）。')

add_body_text(doc,
    '如图22所示，程序运行结束后，右侧结果区显示各设备类型的6项统计表格，'
    '包括最大值(max)、最小值(min)、平均值(mean)、均方根值(rms)、'
    '方差(var)和标准差(std)。中部折线图区以不同颜色曲线展示各设备的蒸汽压力时序对比，'
    '中部箱线图区展示各设备应力分布的箱线图（含中位数、四分位数和异常值），'
    '便于直观比较不同设备类型的应力水平和波动程度。')

add_image_placeholder(doc, 'placeholder_fig22.png', width_cm=14.65, height_cm=8.2)
add_figure_caption(doc, '图22 蒸汽压力分析计算结果')

add_empty_centered(doc)

# ============================================================
# 3.2 蒸汽流量分析
# ============================================================
add_section_title(doc, '3.2 蒸汽流量分析')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，首先点击"选择数据"按钮在电脑硬盘中选择蒸汽流量数据文件。'
    '数据格式与蒸汽压力模块相同，为三列结构（设备标签、数值、单位），'
    '单位为"m³/h"。系统通过单位模糊匹配（contains(unit, "m3") '
    '或 contains(unit, "m³")）自动筛选蒸汽流量相关数据。'
    '同一设备标签的多条时序记录将被自动归入同一组进行计算。')

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '通过点击"蒸汽流量 m³/h"按钮，程序开始运行并执行与蒸汽压力模块相同的7层处理流程'
    '（文件校验→读取→解析→筛选→分组→统计→渲染）。'
    '如图23所示，程序运行结束后，在右侧结果区显示各设备类型的6项统计指标表格。'
    '折线图和箱线图分别展示各设备蒸汽流量的时序变化趋势和分布特征。'
    '用户可切换设备选择或调整Y轴范围来聚焦特定设备的分析结果。')

add_image_placeholder(doc, 'placeholder_fig23.png', width_cm=14.65, height_cm=8.2)
add_figure_caption(doc, '图23 蒸汽流量分析计算结果')

add_empty_centered(doc)

# ============================================================
# 3.3 蒸汽颗粒物浓度分析
# ============================================================
add_section_title(doc, '3.3 蒸汽颗粒物浓度分析')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，首先点击"选择数据"按钮在电脑硬盘中选择蒸汽颗粒物浓度数据文件。'
    '数据格式为三列结构（设备标签、数值、单位），单位为"%"（百分比）。'
    '系统通过单位模糊匹配自动筛选颗粒物浓度相关数据，'
    '并按设备标签中的类型名进行分组。该模块数据量级通常较小（0~1%范围），'
    '系统会自动适配Y轴范围以合理展示数据。')

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '通过点击"蒸汽颗粒物浓度 %"按钮，程序开始运行并按照标准处理流程执行数据分析。'
    '如图24所示，程序运行结束后显示各设备类型的6项统计结果。'
    '折线图自动根据数据量调整线型和标记（≤50条：粗线+圆点标记；'
    '50~200条：中线+小点标记；>200条：细线无标记），确保图表清晰可读。'
    '箱线图展示颗粒物浓度的中位数、分布范围和离群值，便于识别异常设备。')

add_image_placeholder(doc, 'placeholder_fig24.png', width_cm=14.65, height_cm=8.2)
add_figure_caption(doc, '图24 蒸汽颗粒物浓度分析计算结果')

add_empty_centered(doc)

# ============================================================
# 3.4 盐雾浓度分析
# ============================================================
add_section_title(doc, '3.4 盐雾浓度分析')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，首先点击"选择数据"按钮在电脑硬盘中选择盐雾浓度数据文件。'
    '数据格式为三列结构（设备标签、数值、单位），单位为"mg/m³"。'
    '系统通过单位模糊匹配（contains(unit, "mg/m")）自动筛选盐雾浓度相关数据行。'
    '文件中的单位列允许大小写变体（如"mg/m3"、"mg/m³"），系统均能正确识别匹配。')

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '通过点击"盐雾浓度 mg/m³"按钮，程序开始运行。'
    '系统首先检查文件路径、文件存在性和格式合法性，若出现异常则通过errordlg弹窗提示用户。'
    '数据处理过程中，若数值列存在无效值（NaN），系统自动跳过并在状态栏报告跳过数量。'
    '若设备标签为空，自动归入"未命名"组。'
    '如图25所示，程序运行结束后显示完整的统计分析结果。')

add_image_placeholder(doc, 'placeholder_fig25.png', width_cm=14.65, height_cm=8.2)
add_figure_caption(doc, '图25 盐雾浓度分析计算结果')

add_empty_centered(doc)

# ============================================================
# 3.5 氧浓度分析
# ============================================================
add_section_title(doc, '3.5 氧浓度分析')

add_section_title(doc, '（1）参数输入')

add_body_text(doc,
    '开始使用程序，首先点击"选择数据"按钮在电脑硬盘中选择氧浓度数据文件。'
    '数据格式为三列结构（设备标签、数值、单位），单位为"%"（百分比）。'
    '系统通过单位模糊匹配自动筛选氧浓度相关数据，并按设备标签中的类型名自动分组。'
    '每个参数模块使用独立的数据文件，模块之间互不依赖，'
    '用户可随时切换模块进行分析。')

add_section_title(doc, '（2）程序计算')

add_body_text(doc,
    '通过点击"氧浓度 %"按钮，程序开始运行并执行完整的数据分析流程。'
    '系统具备完善的异常处理机制，共分7层24项异常场景覆盖：'
    '文件路径异常、文件读取异常、数据解析异常、数据筛选异常、统计计算异常、'
    '绘图渲染异常和导出异常，确保系统在任何情况下都不会崩溃，'
    '并通过状态栏或弹窗向用户提供明确的错误信息和排查建议。')

add_body_text(doc,
    '如图26所示，程序运行结束后，右侧结果区显示各设备类型的6项统计指标。'
    '用户可点击"导出图表"按钮将折线图和箱线图导出为PNG（300dpi）或FIG格式文件。'
    '点击"刷新图表"按钮可重新加载数据和刷新显示。'
    '当某设备类型仅含1条记录时，方差和标准差自动填写为NaN（空值），'
    '最大值、最小值、平均值和均方根值正常计算。')

add_image_placeholder(doc, 'placeholder_fig26.png', width_cm=14.65, height_cm=8.2)
add_figure_caption(doc, '图26 氧浓度分析计算结果')

add_empty_centered(doc)

# ============================================================
# 补充说明：异常处理与数据格式
# ============================================================
add_section_title(doc, '3.6 数据格式与异常处理说明')

add_body_text(doc,
    '设备应力综合分析模块支持以下数据文件格式：.xlsx、.xls、.csv、.txt。'
    'CSV文件使用逗号分隔，TXT文件使用制表符（Tab）分隔。'
    '数据文件须包含至少两列（设备标签列和数值列），第三列（单位列）为可选。'
    '若无单位列，系统将跳过单位筛选步骤，将全部数据纳入分析。'
    '每个数据文件内同一参数的数据应保持单位一致。')

add_body_text(doc,
    '系统异常处理遵循三项原则：① 不崩溃——任何异常都不得导致系统崩溃；'
    '② 可恢复——处理后系统回到可操作状态；③ 有提示——用户始终能通过状态栏或弹窗'
    '了解当前状态。常见异常场景包括：文件不存在或格式不支持、数据列数不足、'
    '数值列全部无效、无匹配单位的数据等，系统均会给出明确的errordlg弹窗提示。')

add_body_text(doc,
    '对于设备种类超过50种的大规模分析场景，系统会弹出确认对话框询问用户是否继续。'
    '对于仅含1~2条记录的设备类型，箱线图会静默跳过（需至少3个数据点），'
    '但统计表格中仍会正常显示其6项指标（n=1时方差和标准差为NaN）。')

# ============================================================
# 保存
# ============================================================
doc.save(OUT_PATH)
print(f"\n{'=' * 60}")
print(f"文档已生成：{OUT_PATH}")
print(f"文件大小：{os.path.getsize(OUT_PATH) / 1024:.1f} KB")
print(f"{'=' * 60}")
